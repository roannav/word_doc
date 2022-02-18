#!/usr/bin/env python3

import docx

def get_text_from_word_doc( filename, double_space_between_paragraph=True):
    doc = docx.Document(filename)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    paragraph_separator = '\n'
    if double_space_between_paragraph:
        paragraph_separator = '\n\n'
    return paragraph_separator.join(text)


def print_info_about_word_doc( filename):
    doc = docx.Document(filename)
    print(f"\n\n{60*'_'}\nInfo about {filename}")
    num_paragraphs = len(doc.paragraphs)
    print(f"Has {num_paragraphs} paragraphs")

    for i in range(num_paragraphs):
        print(f"\nParagraph #{i}")
        print(doc.paragraphs[i].text)

        num_runs = len(doc.paragraphs[i].runs)
        print(f"\tHas {num_runs} runs")
        # There is always 1+ runs per paragraph
        for r in range(num_runs):
            print(f"\n\tParagraph #{i},   run #{r}")
            print(f"\t\thas style: {doc.paragraphs[i].runs[r].style}")
            print(f"\t\t{doc.paragraphs[i].runs[r].text}")



def run_tests():
    print(f"\n\n{60*'_'}\nTesting get_text_from_word_doc()\n")
    print(get_text_from_word_doc('food_replicator.docx'))

    print(f"\n\n{60*'_'}\nTesting get_text_from_word_doc()\n")
    print(get_text_from_word_doc('food_replicator.docx', False))

    print_info_about_word_doc('food_replicator.docx')


run_tests()
