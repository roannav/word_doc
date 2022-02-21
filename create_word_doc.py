# Create a sample .docx document
# with headings, paragraphs, runs, line breaks, pictures,
# Each run signifies that the style is changing.
# Each run has a different style, than the one before it.
#
# Styles can affect a paragraph or a run.
# The font can also be modified directly.

from docx import Document, text
from docx.shared import Pt, Inches

#print(dir(text.run.WD_BREAK))
print(dir(text.run))

doc = Document()
doc.add_heading('The Late Show with Stephen Colbert')
doc.add_picture('img/The_Late_Show_with_Stephen_Colbert.png', width=Inches(3))

p = doc.add_paragraph('Go to ')
p.add_run('YouTube channel: ').bold = True
p.add_run('https://www.youtube.com/channel/UCMtFAi84ehTSYSE9XoHefig ').add_break()
p.add_run('This channel has 8.7 Million Subscribers')

doc.add_heading('About the show', level=2)
doc.add_paragraph('The Late Show with Stephen Colbert is the premier late ' +
'night talk show on CBS, airing at 11:35pm EST, streaming online via ' +
'Paramount+, and delivered to the International Space Station on a USB ' +
'drive taped to a weather balloon. ', style='Quote')

doc.add_paragraph('Every night, viewers can expect: ' +
'Comedy, humor, funny moments, witty interviews, celebrities, famous ' +
'people, movie stars, bits, humorous celebrities doing bits, funny celebs, big group photos of every star from Hollywood, even the reclusive ones, plus also jokes.', style='Intense Quote')




p = doc.add_paragraph('owned ')

b_run = p.add_run('by ')
b_run.font.small_caps = True

c_run = p.add_run('CBS ', 'Emphasis')
c_run.font.double_strike = True
c_run.font.size = Pt(20)

v_run = p.add_run('Viacom ')
v_run.font.strike = True
v_run.font.size = Pt(25)

pp_run = p.add_run('Paramount Plus')
font = pp_run.font
font.name = 'Calibri'
font.size = Pt(30)
font.shadow = True
font.emboss = True
# This is the BEST way to set font properties




p = doc.add_paragraph('owned by ')
#p.add_run('CBS ').double_strike = True   # doesn't work: no effect
p.add_run('CBS ')
p.add_run('Viacom ')
p.add_run('Paramount Plus')
# p.runs[1].style = 'double_strike'   # doesn't work: compile error
p.runs[0].italic = True               # works
p.runs[1].double_strike = True        # does not work
p.runs[2].bold = True                 # works
p.runs[3].all_caps = True             # does not work




# without any styles on the run objects
p = doc.add_paragraph('You may also like The Late Late Show with James')

# with different styles on each run object
# Results:  only underline works!  Also rtl makes it look different.
p = doc.add_paragraph('You ')
p.add_run('may ').underline = True
p.add_run('also like ').strike = True
p.add_run('The ').small_caps = True
p.add_run('Late ').shadow = True
p.add_run('Late ').outline = True
p.add_run('Show ').rtl = True
p.add_run('with ').imprint = True
p.add_run('James').emboss = True

doc.save('The_Late_Show_with_Stephen_Colbert.docx')
